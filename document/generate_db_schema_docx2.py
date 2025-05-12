from docx import Document

# Danh sách bảng và mô tả chi tiết
summary_tables = [
    (1, "academic_class_courses", "Liên kết các lớp học với các khóa học cụ thể, cho phép một lớp học có thể học nhiều khóa học khác nhau."),
    (2, "academic_class_instructors", "Liên kết các lớp học với các giảng viên phụ trách, hỗ trợ quản lý giảng viên cho từng lớp."),
    (3, "academic_classes", "Lưu thông tin về các lớp học học thuật: mã lớp, tên lớp, năm học, trạng thái hoạt động."),
    (4, "assignment_submissions", "Lưu trữ bài nộp của học viên cho từng bài tập, bao gồm nội dung, file đính kèm, trạng thái chấm điểm."),
    (5, "assignments", "Quản lý các bài tập của khóa học hoặc lớp học, gồm tiêu đề, mô tả, hạn nộp, điểm tối đa, loại bài tập."),
    (6, "categories", "Danh mục các lĩnh vực/nhóm khóa học, giúp phân loại và tổ chức các khóa học theo chủ đề."),
    (7, "certificates", "Lưu thông tin chứng chỉ hoàn thành khóa học của học viên, gồm số chứng chỉ, ngày cấp, trạng thái."),
    (8, "chatbot_response", "Lưu các câu hỏi và phản hồi tự động của chatbot hỗ trợ học viên hoặc người dùng."),
    (9, "class_schedules", "Quản lý lịch học của từng lớp: thời gian, địa điểm, hình thức học, giảng viên, trạng thái lịch học."),
    (10, "course_lesson_discussions", "Lưu các thảo luận, bình luận của học viên về từng bài học, hỗ trợ hỏi đáp và trao đổi kiến thức."),
    (11, "course_lessons", "Quản lý các bài học trong từng phần của khóa học, gồm tiêu đề, nội dung, video, thứ tự, trạng thái."),
    (12, "course_progress", "Theo dõi tiến độ học tập của học viên trong từng khóa học, từng bài học, trạng thái hoàn thành."),
    (13, "course_sections", "Chia nhỏ khóa học thành các phần (section), giúp tổ chức nội dung học tập hợp lý và dễ theo dõi."),
    (14, "courses", "Lưu thông tin chi tiết về các khóa học: tiêu đề, mô tả, giảng viên, giá, trình độ, trạng thái, ảnh."),
    (15, "documents", "Quản lý tài liệu đính kèm cho từng khóa học: tiêu đề, loại file, đường dẫn, kích thước file."),
    (16, "enrollments", "Lưu thông tin đăng ký khóa học của học viên, trạng thái đăng ký, ngày đăng ký, ngày hoàn thành."),
    (17, "forum_likes", "Lưu lượt thích (like) của người dùng cho các chủ đề/thảo luận trên diễn đàn."),
    (18, "forum_replies", "Lưu các phản hồi, trả lời trong các chủ đề diễn đàn, hỗ trợ phân cấp trả lời (parent_id)."),
    (19, "forums", "Quản lý các chủ đề thảo luận của từng khóa học, gồm tiêu đề, mô tả, trạng thái, người tạo."),
    (20, "instructor_availability", "Lưu lịch rảnh của giảng viên để sắp xếp lịch học, lịch dạy phù hợp với từng giảng viên."),
    (21, "messages", "Lưu tin nhắn giữa các người dùng trong hệ thống (học viên, giảng viên, admin)."),
    (22, "notifications", "Quản lý các thông báo gửi đến người dùng: tiêu đề, nội dung, loại thông báo, trạng thái đã đọc."),
    (23, "payments", "Lưu thông tin thanh toán khóa học: số tiền, phương thức, trạng thái, mã giao dịch, ngày thanh toán."),
    (24, "quiz_attempts", "Lưu thông tin mỗi lần học viên làm bài kiểm tra: thời gian bắt đầu, kết thúc, điểm số, trạng thái."),
    (25, "quiz_options", "Lưu các lựa chọn đáp án cho từng câu hỏi kiểm tra, xác định đáp án đúng/sai."),
    (26, "quiz_questions", "Quản lý các câu hỏi trong bài kiểm tra: nội dung, loại câu hỏi, điểm số, giải thích đáp án."),
    (27, "quiz_responses", "Lưu câu trả lời của học viên cho từng câu hỏi trong mỗi lần làm bài kiểm tra."),
    (28, "quizzes", "Quản lý các bài kiểm tra của khóa học/lớp học: tiêu đề, mô tả, thời gian, điểm đạt, trạng thái."),
    (29, "reviews", "Lưu đánh giá của học viên về khóa học hoặc giảng viên: số sao, nội dung đánh giá, loại đánh giá."),
    (30, "user_admins", "Lưu thông tin chi tiết về các quản trị viên: phòng ban, chức vụ, quyền hạn, liên hệ khẩn cấp."),
    (31, "user_grades", "Lưu điểm số của học viên cho từng bài tập, bài kiểm tra, loại điểm, nhận xét, người chấm điểm."),
    (32, "user_instructors", "Lưu thông tin chi tiết về giảng viên: chuyên môn, học vấn, kinh nghiệm, trạng thái xác minh."),
    (33, "user_students", "Lưu thông tin chi tiết về học viên: ngày sinh, giới tính, học vấn, mục tiêu học tập, thành tích."),
    (34, "user_students_academic", "Liên kết học viên với lớp học học thuật, mã sinh viên, năm học, trạng thái học tập."),
    (35, "users", "Lưu thông tin tài khoản người dùng: tên đăng nhập, email, mật khẩu, vai trò, trạng thái, avatar."),
]

doc = Document()
doc.add_heading('Tổng quan lược đồ cơ sở dữ liệu hệ thống E-Learning', 0)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light List Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'STT'
hdr_cells[1].text = 'Tên Bảng'
hdr_cells[2].text = 'Mô Tả'

for stt, name, desc in summary_tables:
    row_cells = table.add_row().cells
    row_cells[0].text = str(stt)
    row_cells[1].text = name
    row_cells[2].text = desc

doc.save('database_schema_summary.docx')
print('Đã tạo file database_schema_summary.docx thành công!') 