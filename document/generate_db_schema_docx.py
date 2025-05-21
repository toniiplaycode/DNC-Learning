from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Dữ liệu chi tiết cho từng bảng (đã cập nhật theo Structure.sql)
db_schema = [
    {
        "name": "Khóa học lớp học (academic_class_courses)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("class_id", "bigint", "FK", "Tham chiếu academic_classes(id)"),
            ("course_id", "bigint", "FK", "Tham chiếu courses(id)"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Giảng viên lớp học (academic_class_instructors)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("class_id", "bigint", "FK", "Tham chiếu academic_classes(id)"),
            ("instructor_id", "bigint", "FK", "ID của user_instructors"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Lớp học chính quy (academic_classes)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("class_code", "varchar(50)", "UNIQUE", "Mã lớp"),
            ("class_name", "varchar(255)", "", "Tên lớp"),
            ("semester", "varchar(20)", "", "Học kỳ (VD: 20231)"),
            ("status", "enum", "", "Trạng thái (active/completed/cancelled)"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Bài nộp (assignment_submissions)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("assignment_id", "bigint", "FK", "Tham chiếu assignments(id)"),
            ("user_id", "bigint", "FK", "Tham chiếu users(id)"),
            ("submission_text", "text", "", "Nội dung nộp"),
            ("file_url", "varchar(255)", "", "File đính kèm"),
            ("submitted_at", "timestamp", "", "Thời gian nộp"),
            ("status", "enum", "", "Trạng thái (submitted/graded/late/resubmit)"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Bài tập (assignments)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("lesson_id", "bigint", "FK", "Tham chiếu course_lessons(id)"),
            ("academic_class_id", "bigint", "FK", "Lớp học nếu là bài tập học thuật"),
            ("title", "varchar(255)", "", "Tiêu đề"),
            ("description", "text", "", "Mô tả"),
            ("due_date", "timestamp", "", "Hạn nộp"),
            ("max_score", "int", "", "Điểm tối đa"),
            ("file_requirements", "text", "", "Yêu cầu file"),
            ("link_document_required", "text", "", "Yêu cầu tài liệu liên kết"),
            ("assignment_type", "enum", "", "Loại bài tập"),
            ("start_time", "datetime", "", "Bắt đầu"),
            ("end_time", "datetime", "", "Kết thúc"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Danh mục khóa học (categories)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("name", "varchar(100)", "", "Tên danh mục"),
            ("description", "text", "", "Mô tả"),
            ("status", "enum", "", "Trạng thái (active/inactive)"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Chứng chỉ (certificates)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("user_id", "bigint", "FK", "Tham chiếu users(id)"),
            ("course_id", "bigint", "FK", "Tham chiếu courses(id)"),
            ("certificate_number", "varchar(100)", "UNIQUE", "Số chứng chỉ"),
            ("certificate_url", "varchar(255)", "", "Đường dẫn chứng chỉ"),
            ("issue_date", "timestamp", "", "Ngày cấp"),
            ("expiry_date", "timestamp", "", "Ngày hết hạn"),
            ("status", "enum", "", "Trạng thái (active/expired/revoked)"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Phản hồi chatbot (chatbot_response)",
        "fields": [
            ("id", "int", "PK", "ID tự tăng"),
            ("keywords", "json", "", "Từ khóa"),
            ("response", "text", "", "Câu trả lời"),
            ("category", "varchar(50)", "", "Danh mục"),
            ("confidence", "float", "", "Độ tin cậy"),
        ]
    },
    {
        "name": "Thảo luận bài học (course_lesson_discussions)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("lesson_id", "bigint", "FK", "Tham chiếu course_lessons(id)"),
            ("user_id", "bigint", "FK", "Tham chiếu users(id)"),
            ("parent_id", "bigint", "FK", "NULL cho thảo luận chính, ID của thảo luận cha cho phản hồi"),
            ("content", "text", "", "Nội dung thảo luận"),
            ("status", "enum", "", "Trạng thái (active/hidden/locked)"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Bài học (course_lessons)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("section_id", "bigint", "FK", "Tham chiếu course_sections(id)"),
            ("title", "varchar(255)", "", "Tiêu đề"),
            ("content_type", "enum", "", "Loại nội dung (video/slide/txt/docx/pdf/xlsx/quiz/assignment)"),
            ("content_url", "varchar(255)", "", "Đường dẫn nội dung"),
            ("content", "text", "", "Nội dung"),
            ("duration", "int", "", "Thời lượng (phút)"),
            ("order_number", "int", "", "Thứ tự"),
            ("is_free", "tinyint(1)", "", "Bài học miễn phí"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Tiến độ học tập (course_progress)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("user_id", "bigint", "FK", "Tham chiếu users(id)"),
            ("lesson_id", "bigint", "FK", "Tham chiếu course_lessons(id)"),
            ("completed_at", "timestamp", "", "Ngày hoàn thành"),
            ("last_accessed", "timestamp", "", "Truy cập cuối"),
        ]
    },
    {
        "name": "Phần học (course_sections)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("course_id", "bigint", "FK", "Tham chiếu courses(id)"),
            ("title", "varchar(255)", "", "Tiêu đề"),
            ("description", "text", "", "Mô tả"),
            ("order_number", "int", "", "Thứ tự"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Khóa học (courses)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("title", "varchar(255)", "", "Tiêu đề"),
            ("description", "text", "", "Mô tả"),
            ("category_id", "bigint", "FK", "Tham chiếu categories(id)"),
            ("instructor_id", "bigint", "FK", "Tham chiếu user_instructors(id)"),
            ("price", "decimal(10,2)", "", "Giá"),
            ("for", "enum", "", "Đối tượng (student/student_academic/both)"),
            ("level", "enum", "", "Trình độ (beginner/intermediate/advanced)"),
            ("status", "enum", "", "Trạng thái (draft/published/archived)"),
            ("thumbnail_url", "varchar(255)", "", "Ảnh đại diện"),
            ("required", "text", "", "Yêu cầu"),
            ("learned", "text", "", "Kết quả đạt được"),
            ("start_date", "date", "", "Ngày bắt đầu"),
            ("end_date", "date", "", "Ngày kết thúc"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Tài liệu (documents)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("instructor_id", "bigint", "FK", "Tham chiếu user_instructors(id)"),
            ("course_section_id", "bigint", "FK", "Tham chiếu course_sections(id)"),
            ("title", "varchar(255)", "", "Tiêu đề"),
            ("description", "text", "", "Mô tả"),
            ("file_url", "varchar(255)", "", "Đường dẫn file"),
            ("file_type", "enum", "", "Loại file (pdf/slide/code/link/txt/docx/xlsx)"),
            ("upload_date", "timestamp", "", "Ngày tải lên"),
            ("download_count", "int", "", "Số lượt tải"),
            ("status", "enum", "", "Trạng thái (active/archived)"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Đăng ký khóa học (enrollments)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("user_id", "bigint", "FK", "Tham chiếu users(id)"),
            ("course_id", "bigint", "FK", "Tham chiếu courses(id)"),
            ("enrollment_date", "timestamp", "", "Ngày đăng ký"),
            ("status", "enum", "", "Trạng thái (active/completed/dropped)"),
            ("completion_date", "timestamp", "", "Ngày hoàn thành"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Lượt thích diễn đàn (forum_likes)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("forum_id", "bigint", "FK", "Tham chiếu forums(id)"),
            ("user_id", "bigint", "FK", "Tham chiếu users(id)"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Phản hồi diễn đàn (forum_replies)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("forum_id", "bigint", "FK", "Tham chiếu forums(id)"),
            ("user_id", "bigint", "FK", "Tham chiếu users(id)"),
            ("reply_id", "bigint", "FK", "ID phản hồi cha"),
            ("content", "text", "", "Nội dung"),
            ("is_solution", "tinyint(1)", "", "Là giải pháp"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Diễn đàn (forums)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("course_id", "bigint", "FK", "Tham chiếu courses(id)"),
            ("user_id", "bigint", "FK", "Tham chiếu users(id)"),
            ("title", "varchar(255)", "", "Tiêu đề"),
            ("description", "text", "", "Mô tả"),
            ("thumbnail_url", "text", "", "Ảnh đại diện"),
            ("status", "enum", "", "Trạng thái (active/archived/closed)"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Tin nhắn (messages)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("sender_id", "bigint", "FK", "Người gửi"),
            ("receiver_id", "bigint", "FK", "Người nhận"),
            ("message_text", "text", "", "Nội dung"),
            ("reference_link", "text", "", "Liên kết tham chiếu"),
            ("is_read", "tinyint(1)", "", "Đã đọc"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Thông báo (notifications)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("user_id", "bigint", "FK", "Tham chiếu users(id)"),
            ("title", "varchar(255)", "", "Tiêu đề"),
            ("content", "text", "", "Nội dung"),
            ("type", "enum", "", "Loại thông báo (course/assignment/quiz/system/message/schedule)"),
            ("is_read", "tinyint(1)", "", "Đã đọc"),
            ("teaching_schedule_id", "bigint", "FK", "Tham chiếu teaching_schedules(id)"),
            ("notification_time", "datetime", "", "Thời gian thông báo"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Thanh toán (payments)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("user_id", "bigint", "FK", "Tham chiếu users(id)"),
            ("course_id", "bigint", "FK", "Tham chiếu courses(id)"),
            ("amount", "decimal(10,2)", "", "Số tiền"),
            ("payment_method", "enum", "", "Phương thức (credit_card/bank_transfer/e_wallet/zalopay)"),
            ("transaction_id", "varchar(100)", "", "Mã giao dịch"),
            ("status", "enum", "", "Trạng thái (pending/completed/failed/refunded)"),
            ("payment_date", "timestamp", "", "Ngày thanh toán"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Lần làm bài kiểm tra (quiz_attempts)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("user_id", "bigint", "FK", "Tham chiếu users(id)"),
            ("quiz_id", "bigint", "FK", "Tham chiếu quizzes(id)"),
            ("start_time", "timestamp", "", "Bắt đầu"),
            ("end_time", "timestamp", "", "Kết thúc"),
            ("score", "decimal(5,2)", "", "Điểm"),
            ("status", "enum", "", "Trạng thái (in_progress/completed/abandoned)"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Lựa chọn câu hỏi (quiz_options)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("question_id", "bigint", "FK", "Tham chiếu quiz_questions(id)"),
            ("option_text", "text", "", "Nội dung"),
            ("is_correct", "tinyint(1)", "", "Đáp án đúng"),
            ("order_number", "int", "", "Thứ tự"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Câu hỏi kiểm tra (quiz_questions)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("quiz_id", "bigint", "FK", "Tham chiếu quizzes(id)"),
            ("question_text", "text", "", "Nội dung câu hỏi"),
            ("question_type", "enum", "", "Loại câu hỏi (multiple_choice/true_false)"),
            ("correct_explanation", "text", "", "Giải thích đáp án"),
            ("points", "int", "", "Số điểm"),
            ("order_number", "int", "", "Thứ tự"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Câu trả lời kiểm tra (quiz_responses)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("attempt_id", "bigint", "FK", "Tham chiếu quiz_attempts(id)"),
            ("question_id", "bigint", "FK", "Tham chiếu quiz_questions(id)"),
            ("selected_option_id", "bigint", "FK", "Tham chiếu quiz_options(id)"),
            ("score", "decimal(5,2)", "", "Điểm"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Bài kiểm tra (quizzes)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("lesson_id", "bigint", "FK", "Tham chiếu course_lessons(id)"),
            ("academic_class_id", "bigint", "FK", "Lớp học nếu là bài kiểm tra học thuật"),
            ("title", "varchar(255)", "", "Tiêu đề"),
            ("description", "text", "", "Mô tả"),
            ("time_limit", "int", "", "Giới hạn thời gian (phút)"),
            ("passing_score", "int", "", "Điểm đạt"),
            ("attempts_allowed", "int", "", "Số lần làm"),
            ("quiz_type", "enum", "", "Loại bài kiểm tra (practice/homework/midterm/final)"),
            ("show_explanation", "tinyint(1)", "", "Hiện giải thích"),
            ("random", "tinyint(1)", "", "Trộn câu hỏi"),
            ("start_time", "datetime", "", "Bắt đầu"),
            ("end_time", "datetime", "", "Kết thúc"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Đánh giá (reviews)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("user_student_id", "bigint", "FK", "Tham chiếu user_students(id)"),
            ("course_id", "bigint", "FK", "Khóa học liên quan"),
            ("review_type", "enum", "", "Loại đánh giá (instructor/course)"),
            ("rating", "int", "", "Số sao (1-5)"),
            ("review_text", "text", "", "Nội dung"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Điểm danh buổi học (session_attendances)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("schedule_id", "bigint", "FK", "Tham chiếu teaching_schedules(id)"),
            ("student_academic_id", "bigint", "FK", "Tham chiếu user_students_academic(id)"),
            ("status", "enum", "", "Trạng thái (present/absent/late/excused)"),
            ("join_time", "datetime", "", "Thời gian tham gia"),
            ("leave_time", "datetime", "", "Thời gian rời đi"),
            ("duration_minutes", "int", "", "Thời lượng (phút)"),
            ("notes", "text", "", "Ghi chú"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Lịch dạy (teaching_schedules)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("academic_class_id", "bigint", "FK", "Tham chiếu academic_classes(id)"),
            ("academic_class_instructor_id", "bigint", "FK", "Tham chiếu academic_class_instructors(id)"),
            ("academic_class_course_id", "bigint", "FK", "Tham chiếu academic_class_courses(id)"),
            ("title", "varchar(255)", "", "Tiêu đề"),
            ("description", "text", "", "Mô tả"),
            ("start_time", "datetime", "", "Bắt đầu"),
            ("end_time", "datetime", "", "Kết thúc"),
            ("meeting_link", "varchar(255)", "", "Link cuộc họp"),
            ("meeting_id", "varchar(100)", "", "ID cuộc họp"),
            ("meeting_password", "varchar(100)", "", "Mật khẩu cuộc họp"),
            ("status", "enum", "", "Trạng thái (scheduled/completed/in-progress/cancelled)"),
            ("is_recurring", "tinyint(1)", "", "Lặp lại"),
            ("recurring_pattern", "json", "", "Mẫu lặp lại"),
            ("recording_url", "text", "", "URL ghi lại"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Quản trị viên (user_admins)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("user_id", "bigint", "FK UNIQUE", "Tham chiếu users(id)"),
            ("full_name", "varchar(100)", "", "Họ tên"),
            ("department", "varchar(100)", "", "Phòng ban"),
            ("position", "varchar(100)", "", "Chức vụ"),
            ("admin_level", "enum", "", "Cấp quản trị (super_admin/admin/moderator)"),
            ("permissions", "json", "", "Quyền hạn"),
            ("emergency_contact", "varchar(100)", "", "Liên hệ khẩn cấp"),
            ("office_location", "varchar(255)", "", "Văn phòng"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Điểm số (user_grades)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("user_id", "bigint", "FK", "Tham chiếu users(id)"),
            ("graded_by", "bigint", "FK", "ID của giảng viên chấm điểm"),
            ("course_id", "bigint", "FK", "Tham chiếu courses(id)"),
            ("lesson_id", "bigint", "FK", "Tham chiếu course_lessons(id)"),
            ("assignment_submission_id", "bigint", "FK", "Tham chiếu assignment_submissions(id)"),
            ("quiz_attempt_id", "bigint", "FK", "Tham chiếu quiz_attempts(id)"),
            ("grade_type", "enum", "", "Loại điểm (assignment/quiz/midterm/final/participation)"),
            ("score", "decimal(5,2)", "", "Điểm"),
            ("max_score", "decimal(5,2)", "", "Điểm tối đa"),
            ("weight", "decimal(5,2)", "", "Trọng số điểm"),
            ("feedback", "text", "", "Nhận xét"),
            ("graded_at", "timestamp", "", "Ngày chấm"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Giảng viên (user_instructors)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("user_id", "bigint", "FK UNIQUE", "Tham chiếu users(id)"),
            ("full_name", "varchar(100)", "", "Họ tên"),
            ("professional_title", "varchar(100)", "", "Chức danh"),
            ("specialization", "varchar(255)", "", "Chuyên môn"),
            ("education_background", "text", "", "Học vấn"),
            ("teaching_experience", "text", "", "Kinh nghiệm"),
            ("bio", "text", "", "Giới thiệu"),
            ("expertise_areas", "text", "", "Lĩnh vực chuyên môn"),
            ("certificates", "text", "", "Chứng chỉ"),
            ("linkedin_profile", "varchar(255)", "", "LinkedIn"),
            ("website", "varchar(255)", "", "Website"),
            ("payment_info", "json", "", "Thông tin thanh toán"),
            ("verification_status", "enum", "", "Trạng thái xác minh (pending/verified/rejected)"),
            ("verification_documents", "text", "", "Tài liệu xác minh"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Học viên (user_students)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("user_id", "bigint", "FK UNIQUE", "Tham chiếu users(id)"),
            ("full_name", "varchar(100)", "", "Họ tên"),
            ("date_of_birth", "date", "", "Ngày sinh"),
            ("gender", "enum", "", "Giới tính (male/female/other)"),
            ("education_level", "varchar(100)", "", "Trình độ học vấn"),
            ("occupation", "varchar(100)", "", "Nghề nghiệp"),
            ("bio", "text", "", "Giới thiệu"),
            ("interests", "text", "", "Sở thích"),
            ("address", "text", "", "Địa chỉ"),
            ("city", "varchar(100)", "", "Thành phố"),
            ("country", "varchar(100)", "", "Quốc gia"),
            ("learning_goals", "text", "", "Mục tiêu học tập"),
            ("preferred_language", "varchar(50)", "", "Ngôn ngữ"),
            ("notification_preferences", "json", "", "Tùy chọn thông báo"),
            ("total_courses_enrolled", "int", "", "Số khóa đã đăng ký"),
            ("total_courses_completed", "int", "", "Số khóa đã hoàn thành"),
            ("achievement_points", "int", "", "Điểm thành tích"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Sinh viên học thuật (user_students_academic)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("user_id", "bigint", "FK UNIQUE", "Tham chiếu users(id)"),
            ("academic_class_id", "bigint", "FK", "Lớp học thuật"),
            ("student_code", "varchar(50)", "UNIQUE", "Mã sinh viên"),
            ("full_name", "varchar(100)", "", "Họ tên"),
            ("academic_year", "varchar(20)", "", "Khóa học (K65, K66...)"),
            ("status", "enum", "", "Trạng thái (studying/graduated/suspended/dropped)"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    },
    {
        "name": "Người dùng (users)",
        "fields": [
            ("id", "bigint", "PK", "ID tự tăng"),
            ("username", "varchar(50)", "", "Tên đăng nhập"),
            ("email", "varchar(100)", "", "Email"),
            ("phone", "varchar(15)", "", "Số điện thoại"),
            ("password", "varchar(255)", "", "Mật khẩu"),
            ("role", "enum", "", "Vai trò (student/instructor/admin/student_academic/chatbot)"),
            ("status", "enum", "", "Trạng thái (active/inactive/banned)"),
            ("avatar_url", "varchar(255)", "", "Ảnh đại diện"),
            ("two_factor_enabled", "tinyint(1)", "", "Bảo mật 2 lớp"),
            ("two_factor_secret", "varchar(100)", "", "Mã bảo mật 2 lớp"),
            ("social_login_provider", "varchar(50)", "", "Đăng nhập MXH"),
            ("social_login_id", "text", "", "ID MXH"),
            ("last_login", "timestamp", "", "Đăng nhập cuối"),
            ("refresh_token", "text", "", "Refresh token"),
            ("created_at", "timestamp", "", "Ngày tạo"),
            ("updated_at", "timestamp", "", "Ngày cập nhật"),
        ]
    }
]

doc = Document()
doc.add_heading('Lược Đồ Cơ Sở Dữ Liệu Hệ Thống E-Learning', 0)

# Create custom style 'bang' if not exists
styles = doc.styles
if 'bang' not in [s.name for s in styles]:
    bang_style = styles.add_style('bang', 1)  # 1 = paragraph style
    bang_style.font.name = 'Times New Roman'
    bang_style.font.size = Pt(13)
    bang_style.font.color.rgb = RGBColor(0, 0, 0)

table_counter = 1

for table in db_schema:
    t = doc.add_table(rows=1, cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    t.style = 'Table Grid'
    
    # Set column widths
    t.columns[0].width = Inches(0.2)  # STT column width
    t.columns[1].width = Inches(2.0)  # Tên thuộc tính
    t.columns[2].width = Inches(1.3)  # Kiểu dữ liệu
    t.columns[3].width = Inches(1.0)  # Khóa
    t.columns[4].width = Inches(3.0)  # Mô tả
    
    # Add right margin to the table
    paragraph = doc.add_paragraph()
    paragraph._p.addnext(t._element)
    doc.paragraphs[-1]._element.getparent().remove(paragraph._element)
    
    hdr_cells = t.rows[0].cells
    hdr_cells[0].text = 'STT'
    hdr_cells[1].text = 'Tên thuộc tính'
    hdr_cells[2].text = 'Kiểu dữ liệu'
    hdr_cells[3].text = 'Khóa'
    hdr_cells[4].text = 'Mô tả'
    
    for idx, field in enumerate(table['fields'], 1):
        row_cells = t.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = field[0]
        row_cells[2].text = field[1]
        row_cells[3].text = field[2]
        row_cells[4].text = field[3]
    
    # Add custom heading 'bang' below the table, centered
    p = doc.add_paragraph(f"Bảng 4.6.{table_counter} {table['name']}", style='bang')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_counter += 1
    doc.add_paragraph()

doc.save('database_schema.docx')
print("Đã tạo file database_schema.docx thành công!")