import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Danh sách các thư mục và tiêu đề
img_groups = [
    ('images/sinhvien', '5.1', 'Giao diện phía sinh viên'),
    ('images/hocvien', '5.2', 'Giao diện phía học viên'),
    ('images/giangvien', '5.3', 'Giao diện phía giảng viên'),
    ('images/quantrivien', '5.4', 'Giao diện phía quản trị viên')
]

# Mô tả chức năng cho từng hình
img_descriptions = {
    # -----------------------------Nhóm sinh viên-----------------------------
    'Trang chủ.png': 'Màn hình chính của hệ thống, cung cấp cái nhìn tổng quan về toàn bộ hoạt động học tập, thông báo mới nhất, các khóa học nổi bật và các chức năng truy cập nhanh. Người dùng có thể dễ dàng theo dõi tiến độ học tập, nhận thông báo quan trọng, truy cập nhanh vào các lớp học, bài tập, diễn đàn hoặc các tài liệu học tập. Đây là điểm khởi đầu giúp sinh viên định hướng các hoạt động tiếp theo một cách thuận tiện và hiệu quả.',
    'Chương trình đạo tạo của sinh viên.png': 'Màn hình hiển thị chương trình đào tạo của sinh viên, bao gồm các học kỳ và môn học. Sinh viên có thể xem thông tin chi tiết về từng môn học, bao gồm tên môn học, mã môn học, số tín chỉ, giảng viên phụ trách, thời gian học tập. Sinh viên có thể theo dõi tiến độ học tập, xem điểm số từng môn học và nhận thông báo về các sự kiện học tập sắp tới. Chức năng này giúp sinh viên quản lý lịch học một cách hiệu quả và đảm bảo không bỏ lỡ các hoạt động học tập quan trọng.', 
    'Trang cá nhân tiến độ học.png': 'Màn hình giúp sinh viên theo dõi tiến độ học tập, hiển thị tỷ lệ hoàn thành.',
    'Trang cá nhân bảng điểm.png': 'Màn hình tổng hợp kết quả học tập, gồm điểm bài tập, kiểm tra, trắc nghiệm, và điểm tổng kết khóa học.',
    'Tìm kiếm thông tin.png': 'Màn hình tìm kiếm toàn bộ nội dung trong hệ thống, bao gồm khóa học, bài học, tài liệu, giảng viên, diễn đàn và thông báo. sinh viên có thể sử dụng bộ lọc nâng cao để tìm kiếm nhanh và chính xác.',
    'Thông báo.png': 'Màn hình tổng hợp tất cả các thông báo quan trọng từ hệ thống, giảng viên hoặc quản trị viên.',
    'Thông báo mail.png': 'Đồng thời tích hợp gửi thông báo qua email, cho phép xem chi tiết, đánh dấu đã đọc, lọc theo loại (học tập, tài chính, sự kiện), nhận thông báo đẩy và email mới. Chức năng duy trì kết nối hiệu quả giữa hệ thống và người dùng.',
    'Tham gia học trực tuyến.png': 'Màn hình cho phép học viên tham gia lớp học trực tuyến qua Google Meet, tích hợp điểm danh tự động và theo dõi thời gian. Học viên kết nối, xem lịch học, hệ thống ghi nhận hiện diện và thời lượng, đảm bảo quản lý học tập hiệu quả.',
    'Nộp bài tập.png': 'Màn hình cho phép sinh viên nộp bài tập trực tuyến, tải lên file, nhập nội dung trả lời, kiểm tra thời hạn nộp và nhận phản hồi từ giảng viên. Ngoài ra, học viên có thể xem lại các bài đã nộp, nhận thông báo khi có điểm hoặc nhận xét. Chức năng này giúp quá trình nộp và chấm bài tập trở nên minh bạch, thuận tiện và tiết kiệm thời gian.',
    'Làm bài trắc nghiệm.png': 'Màn hình cho phép sinh viên thực hiện các bài kiểm tra trắc nghiệm trực tuyến với nhiều dạng câu hỏi khác nhau. Học viên có thể xem thời gian làm bài, chọn đáp án, nộp bài và nhận kết quả ngay sau khi hoàn thành. Ngoài ra, hệ thống còn hỗ trợ lưu tạm thời, xem lại đáp án và thống kê kết quả. Chức năng này giúp đánh giá năng lực học viên một cách khách quan, nhanh chóng và hiệu quả.',
    'Kết quả bài tập.png': 'Màn hình tổng hợp kết quả các bài tập đã nộp, điểm số, nhận xét của giảng viên và trạng thái hoàn thành. sinh viên có thể xem lại bài đã nộp, nhận phản hồi chi tiết và cải thiện kết quả học tập.',
    'Hiển thị điểm và đáp án.png': 'Màn hình hiển thị chi tiết điểm số và đáp án đúng của các bài kiểm tra trắc nghiệm. sinh viên có thể so sánh đáp án của mình với đáp án đúng, nhận nhận xét và rút kinh nghiệm cho các lần kiểm tra sau.',
    'Chi tiết nội dung học video.png': 'Màn hình phát video bài giảng với các tính năng điều chỉnh tốc độ, tua nhanh/chậm, bật/tắt phụ đề. Học viên có thể học mọi lúc, mọi nơi và chủ động kiểm soát quá trình học.',
    'Chi tiết nội dung học văn bản.png': 'Màn hình trình bày nội dung bài học dạng văn bản, hỗ trợ định dạng phong phú, chèn hình ảnh, bảng biểu và liên kết. Học viên có thể đọc và tìm kiếm thông tin nhanh chóng.',
    'Chi tiết nội dung học slide.png': 'Màn hình trình chiếu bài giảng dạng slide, hỗ trợ chuyển trang, phóng to/thu nhỏ, ghi chú và tải về slide. Học viên có thể học tập trực quan, dễ tiếp thu kiến thức.',
    'Chi tiết nội dung học làm trắc nghiệm.png': 'Màn hình cung cấp bài kiểm tra trắc nghiệm trong nội dung học, với nhiều dạng câu hỏi, tính năng lưu tạm, xem lại đáp án và nhận kết quả ngay sau khi nộp. Học viên có thể luyện tập, kiểm tra kiến thức và nhận phản hồi tức thì.',
    'Chi tiết nội dung học làm bài tập.png': 'Màn hình hiển thị chi tiết một bài tập trong nội dung học, bao gồm yêu cầu, hướng dẫn, tài liệu tham khảo và nút nộp bài. Học viên có thể đọc kỹ yêu cầu, tải tài liệu, làm bài và nộp trực tiếp trên hệ thống.',
    'Chi tiết diễn đàn.png': 'Màn hình hiển thị chi tiết một chủ đề thảo luận trong diễn đàn, bao gồm nội dung chủ đề, các bình luận, phản hồi từ học viên, sinh viên và giảng viên. Người dùng có thể đăng bài mới, trả lời, trích dẫn, và nhận thông báo khi có phản hồi mới. Chức năng này tạo môi trường trao đổi học thuật sôi nổi, giúp sinh viên giải đáp thắc mắc và học hỏi lẫn nhau.',
    'Hộp thoại chat nhóm.png':'Tính năng chat nhóm lớp cho phép giao tiếp giữa các thành viên trong lớp học, chia sẻ thông tin và tài liệu học tập. Người dùng có thể gửi file (tài liệu, hình ảnh, video) và sử dụng emoji để tăng tính tương tác và sinh động trong cuộc trò chuyện.',
    'Hộp thoại chat.png': 'Màn hình chat cho phép đổi trực tiếp với người dùng (học viên, giảng viên, quản trị viên), hỗ trợ giải đáp thắc mắc, xử lý sự cố, gửi thông báo khẩn cấp và lưu trữ lịch sử trò chuyện',
    'Giao diện chatbot.png': 'Màn hình trò chuyện với chatbot hỗ trợ học tập, cho phép sinh viên đặt câu hỏi, tìm kiếm tài liệu, nhận hướng dẫn sử dụng hệ thống hoặc giải đáp thắc mắc về bài học.',
    'Các bài tập và trắc nghiệm.png': 'Màn hình tổng hợp tất cả các bài tập và bài kiểm tra trắc nghiệm cần hoàn thành. Sinh viên có thể theo dõi hạn nộp, trạng thái hoàn thành và truy cập nhanh vào từng bài. Nếu là bài nghiệm sẽ có thời gian bắt đầu và kết thúc, sinh viên cần thực hiện theo đúng thời gian. ',
    # -----------------------------Nhóm học viên-----------------------------
    'Chi tiết khóa học.png': 'Màn hình hiển thị chi tiết khóa học: mô tả, mục tiêu, nội dung, bài tập, trắc nghiệm, tài liệu, và thông tin giảng viên, xem đánh giá khóa học',
    'Đăng ký.png': 'Màn hình đăng ký cho học viên nhập thông tin cá nhân, email, số điện thoại, mật khẩu, và đăng ký qua Google. Sau thành công, học viên truy cập học tập, đảm bảo minh bạch, tiện lợi, và mở rộng người dùng.',
    'Đăng nhập.png': 'Màn hình đăng nhập cho học viên  nhập tên, mật khẩu, hoặc đăng nhập qua Google để truy cập khóa học, bài tập, diễn đàn, và tiến độ. Đảm bảo an toàn, bảo mật, và phân quyền theo vai trò.',
    'Danh sách giảng viên.png': 'Màn hình hiển thị danh sách tất cả giảng viên trong hệ thống, cho phép học viên xem thông tin, tìm kiếm theo chuyên ngành, liên hệ hoặc gửi phản hồi.',
    'Danh sách khóa học.png': 'Màn hình tổng hợp tất cả các khóa học trong hệ thống, hỗ trợ tìm kiếm, lọc theo chuyên ngành, trạng thái, giảng viên phụ trách và đăng ký nhanh.',
    'Thanh toán ZaloPay.png': 'Màn hình hỗ trợ học viên thanh toán học phí qua ZaloPay, cho phép quét mã QR, kiểm tra giao dịch, nhận hóa đơn điện tử và xem lịch sử. Chức năng đảm bảo thanh toán nhanh, an toàn, minh bạch.',
    'Trang cá nhân chứng chỉ.png': 'Màn hình hiển thị danh sách chứng chỉ đạt được, cho phép học viên tải về, xem chi tiết, và chia sẻ lên mạng xã hội. Chức năng ghi nhận và khích lệ thành tích học tập.',
    'Trang cá nhân thanh toán.png': 'Màn hình hiển thị chi tiết thanh toán, lịch sử giao dịch, trạng thái học phí, và các khoản phí. Học viên kiểm tra, tải hóa đơn, thanh toán trực tuyến, nhận thông báo thay đổi tài chính, giúp quản lý tài chính học tập hiệu quả.',
    # -----------------------------Nhóm giảng viên-----------------------------
    'Khoa, ngành, chương trình đào tạo của giảng viên.png': 'Màn hình hiển thị thông tin chi tiết về khoa, ngành và chương trình đào tạo mà giảng viên phụ trách. Giảng viên có thể xem danh sách các khoa, ngành học, chương trình đào tạo, số lượng sinh viên đang theo học, và các thông tin liên quan đến công tác giảng dạy. Màn hình này giúp giảng viên quản lý và theo dõi các hoạt động đào tạo trong phạm vi trách nhiệm của mình, bao gồm việc phân công giảng dạy, theo dõi tiến độ học tập của sinh viên và đảm bảo chất lượng đào tạo theo từng chương trình.',
    'Quản lý khóa học.png': 'Đây là màn hình trung tâm cho phép giảng viên hoặc giảng viên quản lý toàn bộ các khóa học trên hệ thống. Người dùng có thể xem danh sách các khóa học hiện có, tìm kiếm, lọc theo chuyên ngành, trạng thái hoặc giảng viên phụ trách. Ngoài ra, chức năng này còn hỗ trợ thêm mới khóa học, chỉnh sửa thông tin chi tiết (tên, mô tả, học phí, thời lượng, giảng viên phụ trách), xóa hoặc tạm ngưng khóa học. Người dùng cũng có thể truy cập vào từng khóa học để quản lý nội dung bài học, tài liệu, bài tập, trắc nghiệm, cũng như theo dõi số lượng học viên đăng ký và tiến độ học tập của từng lớp. Chức năng này giúp đảm bảo việc tổ chức, vận hành và cập nhật các khóa học được thực hiện hiệu quả, đáp ứng nhu cầu đào tạo đa dạng.',
    'Quản lý thông báo.png': 'Màn hình cho phép giảng viên quản lý các thông báo đến học viên và sinh viên trong lớp hoặc toàn bộ hệ thống. Giảng viên có thể phân loại thông báo theo chủ đề, theo dõi trạng thái đã đọc/chưa đọc. Chức năng này giúp đảm bảo thông tin quan trọng được truyền tải kịp thời, tăng cường tương tác giữa giảng viên và học viên.',
    'Quản lý thống kê.png': 'Màn hình này cung cấp các báo cáo, thống kê tổng hợp về hoạt động học tập, giảng dạy và vận hành hệ thống. Người dùng có thể xem thống kê số lượng học viên, giảng viên, khóa học, tỷ lệ hoàn thành khóa học, điểm trung bình, số lượng chứng chỉ đã cấp, doanh thu từ học phí, v.v. Ngoài ra, chức năng còn hỗ trợ lọc, xuất báo cáo theo từng khoảng thời gian, lớp học, khóa học hoặc giảng viên. Thông tin thống kê giúp nhà quản lý đưa ra quyết định kịp thời, tối ưu hóa hoạt động đào tạo và nâng cao chất lượng dịch vụ.',
    'Quản lý diễn đàn (hộp thoại).png': 'Màn hình hộp thoại cho phép giảng viên tạo mới, chỉnh sửa hoặc xóa các chủ đề thảo luận, thiết lập quyền truy cập, đính kèm tài liệu, hình ảnh. Chức năng này giúp quản lý diễn đàn hiệu quả, tạo môi trường học tập tương tác, sáng tạo.',
    'Quản lý diễn đàn.png': 'Màn hình tổng quan về quản lý diễn đàn, cho phép giảng viên xem danh sách tất cả các diễn đàn, chủ đề, số lượng bài viết, bình luận, mức độ tương tác của học viên và thống kê hoạt động thảo luận. Giảng viên có thể tìm kiếm, lọc, truy cập nhanh vào từng diễn đàn để quản lý nội dung.',
    'Quản lý đánh giá.png': 'Màn hình cho phép giảng viên xem, tổng hợp, phân tích và phản hồi các đánh giá từ học viên về chất lượng đào tạo, giảng viên, khóa học, tài liệu. Giảng viên có thể lọc, xuất báo cáo đánh giá, phát hiện các vấn đề nổi bật và đề xuất cải tiến chất lượng đào tạo.',
    'Quản lý điểm danh.png': 'Màn hình cho phép giảng viên điểm danh học viên trong từng buổi học, theo dõi lịch sử điểm danh, xuất báo cáo chuyên cần, gửi cảnh báo cho học viên vắng mặt nhiều lần. Chức năng này giúp nâng cao ý thức học tập, đảm bảo sự tham gia đầy đủ của học viên.',
    'Quản lý lịch dạy (hộp thoại).png': 'Màn hình hộp thoại cho phép giảng viên tạo mới, chỉnh sửa, xóa các buổi học, thiết lập lịch dạy, tạo link Google Meet, gửi thông báo nhắc lịch cho học viên.',
    'Quản lý lịch dạy.png': 'Màn hình tổng quan về lịch dạy của giảng viên, hiển thị chi tiết các buổi học, lớp học, thời gian, nội dung giảng dạy. Giảng viên có thể xem, lọc, tìm kiếm lịch dạy theo tuần, tháng, xuất lịch ra file và nhận thông báo nhắc lịch.',
    'Hộp thoại chat.png': 'Màn hình chat cho phép giảng viên trao đổi trực tiếp với người dùng (học viên, sinh viên và quản trị viên), hỗ trợ giải đáp thắc mắc, xử lý sự cố, và lưu trữ lịch sử trò chuyện. Tính năng chat nhóm lớp cho phép giao tiếp giữa giảng viên và các thành viên trong lớp học, chia sẻ thông tin và tài liệu học tập. Người dùng có thể gửi file (tài liệu, hình ảnh, video) và sử dụng emoji để tăng tính tương tác và sinh động trong cuộc trò chuyện.',
    'Quản lý bài tập (thêm bài cho lớp học thuật).png': 'Màn hình cho phép giảng viên tạo mới bài tập cho từng lớp học thuật, thiết lập yêu cầu, thời hạn nộp, đính kèm tài liệu hướng dẫn, phân loại bài tập theo mức độ khó, chủ đề. Giảng viên có thể gửi thông báo giao bài tập đến học viên.',
    'Quản lý bài trắc nghiệm (thêm bài cho lớp học thuật).png': 'Màn hình cho phép giảng viên tạo mới bài trắc nghiệm cho lớp học thuật (thêm thủ công, tải file trắc nghiệm có sẵn các câu hỏi đúng theo mẫu định sẵn, hoặc sử dụng AI để tự động tạo trắc nghiệm dựa trên tài liệu tải lên), thiết lập câu hỏi, đáp án, thời gian làm bài, phân loại theo chủ đề, mức độ. Hệ thống AI hỗ trợ tạo số lượng câu hỏi trắc nghiệm theo yêu cầu, đồng thời đưa ra số câu hỏi tối đa có thể tạo dựa trên nội dung và độ dài của tài liệu tải lên. Lưu ý, tài liệu không nên quá ngắn để đảm bảo đủ dữ liệu tạo câu hỏi. Giảng viên có thể xem trước bài trắc nghiệm, gửi thông báo cho học viên.',
    'Quản lý bài trắc nghiệm (xem bài làm).png': 'Màn hình cho phép giảng viên xem chi tiết bài làm và điểm đạt được trắc nghiệm của từng học viên hoặc sinh viên.',
    'Quản lý bài trắc nghiệm.png': 'Màn hình tổng quan về quản lý bài trắc nghiệm, cho phép giảng viên tạo mới, chỉnh sửa, xóa, phân loại, xuất kết quả, thống kê điểm số và gửi thông báo cho học viên.',
    'Quản lý bài tập (chấm điểm).png': 'Màn hình cho phép giảng viên chấm điểm bài tập, nhập nhận xét chi tiết, gửi phản hồi cho từng học viên, xuất bảng điểm và thống kê kết quả học tập.',
    'Quản lý bài tập.png': 'Màn hình này cung cấp đầy đủ các nghiệp vụ liên quan đến quản lý bài tập cho từng lớp học hoặc khóa học. Giảng viên có thể tạo mới bài tập, thiết lập thời hạn nộp, mô tả yêu cầu, đính kèm tài liệu hướng dẫn. Ngoài ra, chức năng còn cho phép xem danh sách bài tập đã giao, chỉnh sửa hoặc xóa bài tập, theo dõi số lượng học viên đã nộp bài, chấm điểm trực tiếp trên hệ thống, nhận xét và trả bài cho học viên. Đối với từng bài tập, giảng viên có thể xem chi tiết từng bài nộp, tải về file đính kèm, và gửi phản hồi cá nhân hóa cho từng học viên. Chức năng này giúp nâng cao hiệu quả quản lý, đánh giá và hỗ trợ học viên trong quá trình học tập.',
    'Quản lý lớp học thuật (thêm sinh viên).png': 'Màn hình cho phép giảng viên thêm sinh viên vào lớp học thuật, tìm kiếm, lọc, gửi thông báo mời tham gia lớp.',
    'Quản lý lớp học thuật (thêm khóa học).png': 'Màn hình cho phép giảng viên thêm khóa học vào lớp học thuật để các sinh viên có thể tham gia khóa học. Hệ thống hiển thị các khóa học đã được thêm sẵn bởi chương trình đào tạo, đồng thời cho phép giảng viên xem được các khóa học do các giảng viên khác đã thêm vào.',
    'Quản lý lớp học thuật (hộp thoại).png': 'Màn hình hộp thoại cho phép giảng viên quản lý thông tin chi tiết của lớp học, chỉnh sửa tên lớp, mô tả, danh sách sinh viên.',
    'Quản lý lớp học thuật.png': 'Màn hình tổng quan về quản lý lớp học thuật, cho phép giảng viên xem, tìm kiếm, lọc, xuất báo cáo danh sách lớp, tiến độ học tập, gửi thông báo cho lớp.',
    'Quản lý học sinh sinh viên.png': 'Màn hình tổng quan cho phép giảng viên quản lý toàn bộ học sinh/sinh viên trong các lớp học phụ trách. Giảng viên có thể xem danh sách học viên, tìm kiếm, lọc theo lớp học, khóa học hoặc trạng thái học tập. Ngoài ra, giảng viên có thể truy cập vào thông tin chi tiết của từng học viên để xem tiến độ học tập, điểm số, lịch sử tham gia các hoạt động học tập và các cảnh báo học vụ nếu có. Chức năng này giúp giảng viên theo dõi và hỗ trợ học viên một cách hiệu quả.',
    'Quản lý học sinh sinh viên (cảnh báo học vụ).png': 'Màn hình cho phép giảng viên tạo và quản lý các cảnh báo học vụ cho học viên có kết quả học tập yếu hoặc vắng mặt nhiều. Giảng viên có thể thiết lập các tiêu chí cảnh báo, gửi mail thông báo đến học viên, theo dõi quá trình cải thiện và đánh giá hiệu quả của các biện pháp hỗ trợ. Chức năng này giúp phát hiện sớm và can thiệp kịp thời để nâng cao chất lượng học tập.',
    'Quản lý học sinh sinh viên (bảng điểm).png': 'Màn hình cho phép giảng viên quản lý và theo dõi bảng điểm của học viên trong các lớp học.',
    'Quản lý học viên/sinh viên 2.png': 'Màn hình bổ sung cho quản lý học viên/sinh viên, hỗ trợ các nghiệp vụ nâng cao như phân tích học lực, đề xuất hỗ trợ cá nhân.',
    'Quản lý học viên/sinh viên.png': 'Màn hình này cho phép người dùng (giảng viên hoặc giảng viên) thực hiện đầy đủ các nghiệp vụ liên quan đến quản lý học viên/sinh viên trong hệ thống. Người dùng có thể xem danh sách học viên, tìm kiếm, lọc theo lớp, khóa học hoặc trạng thái học tập. Ngoài ra, màn hình còn hỗ trợ thêm mới học viên, chỉnh sửa thông tin cá nhân, xóa học viên khỏi hệ thống, cũng như truy cập vào chi tiết từng học viên để xem tiến độ học tập, điểm số, lịch sử tham gia các khóa học và các cảnh báo học vụ nếu có. Chức năng này giúp đảm bảo việc quản lý học viên được thực hiện một cách toàn diện, minh bạch và thuận tiện cho các bên liên quan.',
    'Quản lý nội dung khóa học (hộp thoại bài tập).png': 'Màn hình hộp thoại cho phép giảng viên quản lý, thêm mới, chỉnh sửa, xóa bài tập trong nội dung khóa học, đính kèm tài liệu hướng dẫn, phân loại bài tập.',
    'Quản lý nội dung khóa học (hộp thoại trắc nghiệm).png': 'Màn hình hộp thoại cho phép giảng viên quản lý, thêm mới (thêm thủ công, tải file docx có sẵn các câu hỏi trắc nghiệm theo mẫu định sẵn, hoặc sử dụng AI để tự động tạo trắc nghiệm dựa trên tài liệu tải lên), chỉnh sửa, xóa bài trắc nghiệm trong nội dung khóa học, thiết lập câu hỏi, đáp án. Hệ thống AI hỗ trợ tạo số lượng câu hỏi trắc nghiệm theo yêu cầu, đồng thời đưa ra số câu hỏi tối đa có thể tạo dựa trên nội dung và độ dài của tài liệu tải lên. Lưu ý, tài liệu không nên quá ngắn so với số lượng câu hỏi muốn tạo.',
    'Quản lý nội dung khóa học (thêm tài liệu).png': 'Màn hình cho phép giảng viên thêm mới tài liệu học tập vào khóa học, đính kèm file, mô tả chi tiết, phân loại tài liệu.',
    'Quản lý nội dung khóa học (thêm nội dung).png': 'Màn hình cho phép giảng viên thêm mới nội dung bài học, bài tập, trắc nghiệm vào khóa học, thiết lập thứ tự, phân loại nội dung.',
    'Quản lý nội dung khóa học.png': 'Màn hình tổng quan về quản lý nội dung khóa học, cho phép giảng viên xem, tìm kiếm, chỉnh sửa, phân loại.',
    'Quản lý khóa học (hộp thoại).png': 'Màn hình hộp thoại cho phép giảng viên quản lý thông tin chi tiết của khóa học, chỉnh sửa tên, mô tả, lịch học, giảng viên phụ trách.',
    'Đăng nhập giảng viên.png': 'Màn hình đăng nhập dành riêng cho giảng viên, cho phép truy cập các chức năng quản lý lớp học, khóa học, bài tập, trắc nghiệm và trao đổi với học viên. Đảm bảo an toàn, bảo mật và phân quyền rõ ràng cho giảng viên.',
    # -----------------------------Nhóm quản trị viên-----------------------------
    'Khoa, ngành, chương trình đào tạo của quản trị viên.png': 'Màn hình quản lý toàn diện về khoa, ngành và chương trình đào tạo trong hệ thống. Quản trị viên có thể thực hiện tạo mới, xem, chỉnh sửa và xóa các khoa, ngành học, chương trình đào tạo. Màn hình này cho phép quản trị viên thiết lập cấu trúc tổ chức đào tạo, phân công giảng viên cho từng khoa/ngành, quản lý số lượng sinh viên, thiết lập các quy định đào tạo và đảm bảo tính nhất quán trong hệ thống quản lý giáo dục. Chức năng này giúp quản trị viên duy trì và cập nhật thông tin đào tạo một cách hiệu quả và chính xác.',
    'Quản lý đánh giá.png': 'Màn hình cho phép quản trị viên xem, tổng hợp, phân tích và phản hồi các đánh giá từ học viên về chất lượng đào tạo, giảng viên, khóa học, tài liệu. Quản trị viên có thể lọc, xuất báo cáo đánh giá, phát hiện các vấn đề nổi bật và đề xuất cải tiến chất lượng đào tạo.',
    'Quản lý thống kê.png': 'Màn hình này cung cấp các báo cáo, thống kê tổng hợp về hoạt động học tập, giảng dạy và vận hành hệ thống. Người dùng có thể xem thống kê số lượng học viên, giảng viên, khóa học, tỷ lệ hoàn thành khóa học, số lượng chứng chỉ đã cấp, doanh thu từ học phí, v.v. Ngoài ra, chức năng còn hỗ trợ lọc, xuất báo cáo theo từng khoảng thời gian, lớp học, khóa học hoặc giảng viên. Thông tin thống kê giúp nhà quản lý đưa ra quyết định kịp thời, tối ưu hóa hoạt động đào tạo và nâng cao chất lượng dịch vụ.',
    'Quản lý thanh toán.png': 'Màn hình tổng quan về quản lý thanh toán học phí, cho phép quản trị viên theo dõi doanh thu, các khoản thu/chi, xuất báo cáo tài chính, kiểm soát các khoản phí phát sinh và đảm bảo minh bạch tài chính trong hệ thống.',
    'Quản lý lịch dạy các giảng viên.png': 'Màn hình tổng quan cho phép quản trị viên quản lý lịch dạy của tất cả giảng viên trong hệ thống. Quản trị viên có thể xem lịch dạy của từng giảng viên, theo dõi lịch học, lớp học, thời gian giảng dạy và nội dung bài giảng. Chức năng này hỗ trợ việc phân công giảng dạy, kiểm soát tải giảng viên, đảm bảo không có xung đột lịch dạy và tối ưu hóa việc sử dụng nguồn lực giảng viên. Quản trị viên có thể xuất báo cáo lịch dạy, gửi thông báo nhắc lịch cho giảng viên và theo dõi hiệu quả công tác giảng dạy.',
    'Danh mục (hộp thoại).png': 'Màn hình hộp thoại cho phép quản trị viên tạo mới, chỉnh sửa, xóa các danh mục khóa học, phân loại chuyên ngành, thiết lập phân loại, đảm bảo hệ thống danh mục luôn cập nhật, khoa học.',
    'Danh mục.png': 'Màn hình tổng quan về quản lý danh mục khóa học, cho phép xem, tìm kiếm, lọc, chỉnh sửa, xuất báo cáo danh mục, hỗ trợ tổ chức hệ thống khóa học hợp lý.',
    'Hộp thoại chat.png': 'Màn hình chat cho phép quản trị viên trao đổi trực tiếp với người dùng (học viên, sinh viên, giảng viên), hỗ trợ giải đáp thắc mắc, xử lý sự cố, và lưu trữ lịch sử trò chuyện. Chia sẻ thông tin và tài liệu quan trọng. Người dùng có thể gửi file (tài liệu, hình ảnh, video) và sử dụng emoji để tăng tính tương tác và sinh động trong cuộc trò chuyện.',
    'Quản lý lớp học thuật (thêm sinh viên).png': 'Màn hình cho phép quản trị viên thêm sinh viên vào lớp học thuật, tìm kiếm, lọc, gửi thông báo mời tham gia lớp.',
    'Quản lý lớp học thuật (phân công).png': 'Màn hình cho phép quản trị viên phân công giảng viên cho từng lớp học thuật, gửi thông báo phân công.',
    'Quản lý lớp học thuật (hộp thoại).png': 'Màn hình hộp thoại cho phép quản trị viên quản lý thông tin chi tiết của lớp học, chỉnh sửa Mã lớp, tên lớp, học kỳ, giảng viên phụ trách.',
    'Quản lý lớp học thuật.png': 'Màn hình tổng quan về quản lý lớp học thuật, cho phép quản trị viên xem, tìm kiếm, lọc, xuất báo cáo danh sách lớp, tiến độ học tập, phân công giảng dạy và phụ trách cho giảng viên.',
    'Quản lý học viên và sinh viên (xem thông tin).png': 'Màn hình cho phép quản trị viên xem chi tiết thông tin học viên hoặc sinh viên, lịch sử học tập, điểm số, trạng thái học phí, cảnh báo học vụ.',
    'Quản lý học viên và sinh viên.png': 'Đây là màn hình tổng quan cho phép quản trị viên hoặc giảng viên quản lý toàn bộ danh sách học viên trong hệ thống. Người dùng có thể tìm kiếm, lọc, thêm mới, chỉnh sửa hoặc xóa học viên, cũng như truy cập vào chi tiết từng học viên để theo dõi tiến độ học tập, điểm số và các thông tin liên quan. Chức năng này giúp việc quản lý học viên trở nên dễ dàng, hiệu quả và chính xác.',
    'Quản lý giảng viên (truy cập trang giảng viên).png': 'Màn hình cho phép quản trị viên truy cập trang của giảng viên, thực hiện được tất cả các quyền và quản lý của giảng viên.',
    'Quản lý giảng viên (hộp thoại).png': 'Màn hình hộp thoại cho phép quản trị viên quản lý thông tin giảng viên, chỉnh sửa hồ sơ, phân công lớp học.',
    'Quản lý giảng viên.png': 'Màn hình tổng quan về quản lý giảng viên, cho phép xem, tìm kiếm, lọc, xuất báo cáo danh sách giảng viên, phân công giảng dạy.',
    'Quản lý khóa học (xem nội dung khóa học).png': 'Màn hình cho phép quản trị viên xem chi tiết nội dung của từng khóa học, bao gồm bài học, bài tập, trắc nghiệm, tài liệu, giảng viên phụ trách.',
    'Quản lý khóa học (hộp thoại).png': 'Màn hình hộp thoại cho phép quản trị viên quản lý thông tin chi tiết của khóa học, chỉnh sửa tên, mô tả, giảng viên phụ trách.',
    'Quản lý khóa học.png': 'Đây là màn hình trung tâm cho phép quản trị viên hoặc giảng viên quản lý toàn bộ các khóa học trên hệ thống. Người dùng có thể xem danh sách các khóa học hiện có, tìm kiếm, lọc theo chuyên ngành, trạng thái hoặc giảng viên phụ trách. Ngoài ra, chức năng này còn hỗ trợ thêm mới khóa học, chỉnh sửa thông tin chi tiết (tên, mô tả, học phí, thời lượng, giảng viên phụ trách), xóa hoặc tạm ngưng khóa học. Người dùng cũng có thể truy cập vào từng khóa học để quản lý nội dung bài học, tài liệu, bài tập, trắc nghiệm, cũng như theo dõi số lượng học viên đăng ký và tiến độ học tập của từng lớp. Chức năng này giúp đảm bảo việc tổ chức, vận hành và cập nhật các khóa học được thực hiện hiệu quả, đáp ứng nhu cầu đào tạo đa dạng.',
    'Đăng nhập quản trị viên.png': 'Màn hình đăng nhập dành riêng cho quản trị viên hệ thống. Tại đây, quản trị viên nhập tên đăng nhập và mật khẩu để truy cập các chức năng quản lý, cấu hình hệ thống, kiểm soát người dùng, khóa học, tài chính và các nghiệp vụ quan trọng khác. Chức năng này đảm bảo an toàn, bảo mật và phân quyền truy cập rõ ràng giữa các vai trò trong hệ thống.'
}

doc = Document()

# Tạo style 'hinh' nếu chưa có
styles = doc.styles
if 'hinh' not in [s.name for s in styles]:
    hinh_style = styles.add_style('hinh', 1)
    hinh_style.font.name = 'Times New Roman'
    hinh_style.font.size = Pt(12)
    hinh_style.font.bold = True
    hinh_style.font.color.rgb = RGBColor(0, 0, 0)

# Tạo style 'mota' cho đoạn mô tả
if 'mota' not in [s.name for s in styles]:
    mota_style = styles['Normal']
    mota_style.font.name = 'Times New Roman'
    mota_style.font.size = Pt(13)
    mota_style.paragraph_format.first_line_indent = Inches(0.5)  # 1.27cm = 0.5 inches
    mota_style.paragraph_format.space_after = Pt(12)
    mota_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Thêm heading chương 5
doc.add_heading('Chương 5. Giao diện hệ thống', level=0)

# Biến đếm số hình toàn cục cho caption
figure_counter = 0

for folder, group_code, group_title in img_groups:
    if not os.path.exists(folder):
        continue
    img_files = [f for f in os.listdir(folder) if f.lower().endswith('.png')]
    # Sắp xếp logic cho từng nhóm
    if folder.endswith('sinhvien'):
        logical_order = [
            'Trang chủ.png',
            'Chương trình đạo tạo của sinh viên.png',
            'Hộp thoại chat nhóm.png',
            'Tham gia học trực tuyến.png',
            'Các bài tập và trắc nghiệm.png',
            'Nộp bài tập.png',
            'Kết quả bài tập.png',
            'Làm bài trắc nghiệm.png',
            'Hiển thị điểm và đáp án.png',
            'Trang cá nhân tiến độ học.png',
            'Trang cá nhân bảng điểm.png',
        ]
    elif folder.endswith('hocvien'):
        logical_order = [
            'Đăng ký.png',
            'Đăng nhập.png',
            'Hộp thoại chat.png',
            'Giao diện chatbot.png',
            'Danh sách khóa học.png',
            'Thanh toán ZaloPay.png',
            'Chi tiết khóa học.png',
            'Chi tiết nội dung học video.png',
            'Chi tiết nội dung học văn bản.png',
            'Chi tiết nội dung học slide.png',
            'Chi tiết nội dung học làm trắc nghiệm.png',
            'Chi tiết nội dung học làm bài tập.png',
            'Danh sách giảng viên.png',
            'Chi tiết diễn đàn.png',
            'Trang cá nhân chứng chỉ.png',
            'Trang cá nhân thanh toán.png',
            'Tìm kiếm thông tin.png',
            'Thông báo.png',
            'Thông báo mail.png',
        ]
    elif folder.endswith('giangvien'):
        logical_order = [
            'Đăng nhập giảng viên.png',
            'Khoa, ngành, chương trình đào tạo của giảng viên.png',
            'Quản lý khóa học.png',
            'Quản lý khóa học (hộp thoại).png',
            'Quản lý nội dung khóa học.png',
            'Quản lý nội dung khóa học (thêm nội dung).png',
            'Quản lý nội dung khóa học (thêm tài liệu).png',
            'Quản lý nội dung khóa học (hộp thoại bài tập).png',
            'Quản lý nội dung khóa học (hộp thoại trắc nghiệm).png',
            'Quản lý học sinh sinh viên.png',
            'Quản lý học sinh sinh viên (cảnh báo học vụ).png',
            'Quản lý học sinh sinh viên (bảng điểm).png',
            'Quản lý lớp học thuật.png',
            'Quản lý lớp học thuật (hộp thoại).png',
            'Quản lý lớp học thuật (thêm sinh viên).png',
            'Quản lý lớp học thuật (thêm khóa học).png',
            'Quản lý bài tập.png',
            'Quản lý bài tập (thêm bài cho lớp học thuật).png',
            'Quản lý bài tập (chấm điểm).png',
            'Quản lý bài trắc nghiệm.png',
            'Quản lý bài trắc nghiệm (thêm bài cho lớp học thuật).png',
            'Quản lý bài trắc nghiệm (xem bài làm).png',
            'Hộp thoại chat.png',
            'Quản lý lịch dạy.png',
            'Quản lý lịch dạy (hộp thoại).png',
            'Quản lý điểm danh.png',
            'Quản lý diễn đàn.png',
            'Quản lý diễn đàn (hộp thoại).png',
            'Quản lý đánh giá.png',
            'Quản lý thống kê.png',
            'Quản lý thông báo.png',
        ]
    elif folder.endswith('quantrivien'):
        logical_order = [
            'Đăng nhập quản trị viên.png',
            'Khoa, ngành, chương trình đào tạo của quản trị viên.png',
            'Quản lý khóa học.png',
            'Quản lý khóa học (hộp thoại).png',
            'Quản lý khóa học (xem nội dung khóa học).png',
            'Quản lý giảng viên.png',
            'Quản lý giảng viên (hộp thoại).png',
            'Quản lý giảng viên (truy cập trang giảng viên).png',
            'Quản lý học viên và sinh viên.png',
            'Quản lý học viên và sinh viên (xem thông tin).png',
            'Quản lý lớp học thuật.png',
            'Quản lý lớp học thuật (hộp thoại).png',
            'Quản lý lớp học thuật (thêm sinh viên).png',
            'Quản lý lớp học thuật (phân công).png',
            'Hộp thoại chat.png',
            'Danh mục.png',
            'Danh mục (hộp thoại).png',
            'Quản lý thanh toán.png',
            'Quản lý lịch dạy các giảng viên.png',
            'Quản lý thống kê.png',
            'Quản lý đánh giá.png',
        ]
    else:
        logical_order = []
    # Sắp xếp theo logic, các file còn lại thêm vào cuối
    # Lọc bỏ các file đã comment trong logical_order (chỉ lấy các dòng không chứa # sau khi loại bỏ khoảng trắng)
    logical_order_filtered = [f.strip() for f in logical_order if '#' not in f.strip()]
    img_files_sorted = [f for f in logical_order_filtered if f in img_files]
    img_files_sorted += [f for f in img_files if f not in img_files_sorted]
    img_files = img_files_sorted
    # Thêm heading nhóm (5.1, 5.2, 5.3, 5.4) là heading 2
    doc.add_heading(f'{group_code} {group_title}', level=2)
    # Lặp qua từng hình trong nhóm
    for idx, img_file in enumerate(img_files, 1):
        # Kiểm tra mô tả chức năng
        description = img_descriptions.get(img_file, 'Chức năng chưa được mô tả.')
        
        # Bỏ qua các hình có mô tả "Chức năng chưa được mô tả"
        if description == 'Chức năng chưa được mô tả.':
            continue
            
        figure_counter += 1
        name = os.path.splitext(img_file)[0].replace('_', ' ')
        # Thêm heading mục 5.x.y là heading 3
        doc.add_heading(f'{group_code}.{idx} {name.capitalize()}', level=3)
        img_path = os.path.join(folder, img_file)
        # Thêm hình vào đoạn văn bản và căn giữa
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.5))
        # Chỉ thay đổi caption thành đánh số tuần tự
        caption = doc.add_paragraph(f'Hình 5.{figure_counter} {name.capitalize()}', style='hinh')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Thêm mô tả chức năng
        desc_para = doc.add_paragraph(description)
        desc_para.style.font.name = 'Times New Roman'
        desc_para.style.font.size = Pt(13)
        desc_para.paragraph_format.first_line_indent = Inches(0.5)
        desc_para.paragraph_format.space_after = Pt(12)
        desc_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.save('bao_cao_chuong5.docx')
print('Đã tạo file bao_cao_chuong5.docx thành công!')